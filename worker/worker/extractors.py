"""
Document text extraction module.
Supports PDF, DOCX, and TXT formats with optimal extraction strategies.
"""
import os
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
from docx import Document as DocxDocument


class DocumentExtractor:
    """Extract text content from various document formats."""

    @staticmethod
    def extract(file_path: str) -> str:
        """
        Extract text from document based on file extension.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file does not exist
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = Path(file_path).suffix.lower()
        
        if ext == '.pdf':
            return DocumentExtractor._extract_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return DocumentExtractor._extract_docx(file_path)
        elif ext == '.txt':
            return DocumentExtractor._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        """Extract text from PDF using PyMuPDF for optimal performance."""
        text_parts = []
        
        with fitz.open(file_path) as doc:
            for page in doc:
                # Extract text with layout preservation
                text = page.get_text("text")
                if text.strip():
                    text_parts.append(text)
        
        return "\n".join(text_parts)
    
    @staticmethod
    def _extract_docx(file_path: str) -> str:
        """Extract text from DOCX files."""
        doc = DocxDocument(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return "\n".join(text_parts)
    
    @staticmethod
    def _extract_txt(file_path: str) -> str:
        """Extract text from plain text files with encoding detection."""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        # Fallback: read as binary and decode with errors='ignore'
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
